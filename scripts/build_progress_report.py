#!/usr/bin/env python3
"""依「6 大模組 × 子選單」結構產生 Word 進度報告。"""
import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT_DIR = Path("/home/audi/project/DMIS/output_report")
SHOTS = OUT_DIR / "screenshots_full"
DOCX = OUT_DIR / "DMIS_專案進度報告_v5.docx"

ZH = "Microsoft JhengHei"
NAVY = RGBColor(0x1F, 0x3A, 0x68)
GRAY = RGBColor(0x55, 0x55, 0x55)


def set_zh(run, size=11, bold=False, color=None):
    run.font.name = ZH
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), ZH)
    rFonts.set(qn("w:hAnsi"), ZH)
    rFonts.set(qn("w:ascii"), ZH)
    run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, size=11, bold=False, color=None, align=None, indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_zh(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 26, 1: 20, 2: 15, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_zh(run, size=sizes.get(level, 12), bold=True, color=NAVY)


def add_image(doc, name, caption, width_in=6.3):
    path = SHOTS / f"{name}.png"
    if not path.exists():
        add_para(doc, f"[缺圖：{name}]", color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    set_zh(cap.add_run(f"圖：{caption}"), size=9, color=GRAY)


def add_steps(doc, steps):
    for i, s in enumerate(steps, 1):
        add_para(doc, f"{i}. {s}", size=11, indent=0.2)


def add_kv_table(doc, kv):
    t = doc.add_table(rows=len(kv), cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(kv):
        c0, c1 = t.rows[i].cells
        c0.width = Inches(1.6)
        c1.width = Inches(4.5)
        set_zh(c0.paragraphs[0].add_run(k), size=10, bold=True, color=NAVY)
        set_zh(c1.paragraphs[0].add_run(v), size=10)


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "DMIS 合作車行管理系統", size=30, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "操作手冊與進度報告（v5）", size=18, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "六大模組｜逐子選單操作說明｜實機畫面", size=13, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(5):
        doc.add_paragraph()
    add_para(doc, f"產出日期：{datetime.date.today().isoformat()}", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
    add_para(doc, "技術架構：Odoo 16 Community｜PostgreSQL 15｜Docker Compose", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
    doc.add_page_break()


def section_overview(doc):
    add_heading(doc, "1. 專案總覽", level=1)
    add_para(doc,
             "DMIS（Dealer Management Information System）為一套服務於汽機車合作車行"
             "體系的營運管理平台，整合「車行」、「銷售」、「傭金」、「零件」、"
             "「報表分析」與「使用者權限」六大領域，協助總部與合作車行即時掌握第一線業務脈動。"
             "本報告依實際選單逐一說明六大模組底下每個子選單的用途與操作步驟，"
             "並附上實機截圖以利使用者按圖索驥。", size=11)

    add_heading(doc, "1.1 系統登入", level=2)
    add_steps(doc, [
        "於瀏覽器開啟 http://localhost:8069 進入登入頁。",
        "輸入帳號（預設 admin）、密碼（預設 admin）後按「登入」。",
        "登入後將進入主應用程式列表，點選對應應用即可使用。",
    ])
    add_image(doc, "00_login", "Odoo 登入畫面")
    add_image(doc, "00_apps", "登入後的應用程式列表")

    add_heading(doc, "1.2 環境與技術", level=2)
    add_kv_table(doc, [
        ("ERP 框架", "Odoo 16 Community（容器化部署）"),
        ("資料庫", "PostgreSQL 15"),
        ("部署方式", "Docker Compose（odoo / db / nginx 三服務）"),
        ("自製模組", "dms_core / dms_sale / dms_commission / dms_parts / "
                     "dms_report / user_management 等"),
        ("BI 報表整合", "Metabase（http://localhost:3000）一鍵建表與儀表板自動部署，"
                       "提供 P1 ~ P22 等多張銷售 / 車行 / 通路 / 客群儀表板"),
    ])

    add_heading(doc, "1.3 報告涵蓋範圍", level=2)
    add_para(doc, "本報告聚焦六大功能選單，每個子選單均提供「用途說明 + 操作步驟 + 實機畫面」。")
    add_kv_table(doc, [
        ("(2) 車行管理", "車行主檔、品牌、類型、拜訪、行事曆、目的、假日、批次與政府同步"),
        ("(3) 產品管理", "產品頁面（模板）、價格表（SKU 世代）、價目版本、分期規則模板、費用類型、規則掛接"),
        ("(4) 車銷管理", "車輛銷售、Excel 匯入、OrderProcessor 同步紀錄（含重新同步）"),
        ("(5) 銷售分析（Metabase）", "以 Metabase 儀表板呈現總車輛 / 油車 / 電車 / 車行 / 通路 / 客群等多維度報表"),
        ("(6) 佣金管理", "基礎/車行/車種規則、台數獎勵、激勵品項、佣金記錄、核銷、月結"),
        ("(7) 零件管理", "零件主檔、分類、目錄、查詢工具、CSV 匯入"),
        ("(8) 使用者管理系統", "存取群組設定、操作歷程稽核"),
    ])
    doc.add_page_break()


SECTIONS = [
    {
        "title": "2. 車行管理（DMS Core）",
        "intro": "車行管理是 DMIS 的主檔來源，所有銷售、傭金、零件、拜訪都會引用此模組的「車行」為維度。"
                 "底下涵蓋主檔資料、拜訪紀錄與行事曆、節假日設定等實務作業。",
        "items": [
            {"h": "2.1 車行（Dealer）",
             "brief": "管理所有合作車行之主檔，包含名稱、編號、地址、聯絡人、品牌與類型。",
             "steps": ["從左上角選單進入「車行管理 → 車行」。",
                       "在清單畫面可使用上方搜尋列以名稱、品牌、城市快速過濾。",
                       "點擊任一列開啟車行表單可檢視/編輯詳細資料；點「建立」可新增車行。",
                       "編輯完成後按「儲存」，異動會自動寫入操作歷程。"],
             "shots": [("11_dealers_list", "車行清單"), ("11_dealers_form", "車行表單")]},
            {"h": "2.2 品牌（Brand）",
             "brief": "維護車行所屬品牌（如 KYMCO、SYM、YAMAHA 等），供主檔下拉選擇。「品牌」選單僅系統管理員（base.group_system）可見；tree 提供品牌 Logo 縮圖欄位。",
             "steps": ["進入「車行管理 → 品牌」（仅 admin 可見）。",
                       "點「建立」輸入品牌名稱，並於 form 上傳 Logo 圖片後儲存。",
                       "品牌資料會在車行、車型、價格表等多處被引用。"],
             "shots": [("12_brands_list", "品牌清單"), ("12_brands_form", "品牌表單")]},
            {"h": "2.3 車行類型（Store Type）",
             "brief": "定義車行型態（如總代理、經銷、加盟、衛星店）；用於分群統計。",
             "steps": ["進入「車行管理 → 車行類型」。",
                       "新增類型並指定編碼，於車行主檔可選擇套用。"],
             "shots": [("13_storetype_list", "車行類型清單"), ("13_storetype_form", "車行類型表單")]},
            {"h": "2.4 拜訪紀錄（Visit）",
             "brief": "業務人員親訪車行時建立的工作紀錄，含日期、目的、結論、附件。",
             "steps": ["進入「車行管理 → 拜訪紀錄」。",
                       "點「建立」開啟新拜訪表單，依序填入車行、日期、拜訪人、目的。",
                       "在「備註」輸入結論與後續事項；可上傳照片或文件。",
                       "送出後會以草稿/完成/取消三狀態流轉，方便主管追蹤。"],
             "shots": [("14_visit_list", "拜訪紀錄清單"), ("14_visit_form", "拜訪紀錄表單")]},
            {"h": "2.5 拜訪行事曆（Calendar）",
             "brief": "以行事曆方式檢視拜訪安排，便於排程與資源調度。",
             "steps": ["進入「車行管理 → 拜訪行事曆」。",
                       "可切換月/週/日視圖；直接點空白格新增拜訪。"],
             "shots": [("15_visit_calendar", "拜訪行事曆視圖")]},
            {"h": "2.6 拜訪目的類別",
             "brief": "標準化拜訪原因（例：例行訪、新品介紹、催收、稽核）。",
             "steps": ["進入「車行管理 → 拜訪目的類別」並建立常用類別。"],
             "shots": [("16_visit_purpose", "拜訪目的類別清單")]},
            {"h": "2.7 台灣假日設定",
             "brief": "存放國定假日與調整工作日，影響拜訪頻次與工作天計算。",
             "steps": ["進入「車行管理 → 台灣假日設定」。",
                       "可手動新增單筆假日，或使用下一節「同步政府假日」批次匯入。"],
             "shots": [("17_holiday_list", "台灣假日清單")]},
            {"h": "2.8 同步政府假日",
             "brief": "從政府開放資料平台一鍵抓取年度假期到本系統。",
             "steps": ["點選「車行管理 → 同步政府假日」開啟精靈。",
                       "選擇要同步的年度，按「執行」即會抓取並寫入。"],
             "shots": [("18_holiday_sync_wizard", "同步政府假日精靈")]},
            {"h": "2.9 批次建立拜訪",
             "brief": "依條件（區域、車行清單、頻率）一次產生多筆拜訪計畫。",
             "steps": ["進入「車行管理 → 批次建立拜訪」精靈。",
                       "選擇車行範圍、預設拜訪人、起訖日期、間隔。",
                       "預覽後確認執行，系統將大量建立草稿拜訪供後續編輯。"],
             "shots": [("19_visit_bulk_wizard", "批次建立拜訪精靈")]},
        ],
    },
    {
        "title": "3. 產品管理（DMS Product）",
        "intro": "產品管理模組是 015 重建後的独立模組，包含「產品頁面（模板）」、「車型價格表」、「價目版本」、「分期規則模板」、「費用類型」、「規則掛接」等選單。車銷與佣金依賴本模組提供的主檔與價格資料。",
        "items": [
            {"h": "3.1 產品頁面（模板，Product Template）",
             "brief": "產品模板主檔，以「品牌 + 機種 + 型號」為主鍵建立。表單下頁簽提供「產品項」可見下屬車色 / SKU、「價格」可見各價目版本中本模板的牌價 / 現金價 / 分期價。tree 中最左側提供「開啟詳細資料」外部連結按鈕，使用者可在不離開清單的情況下，横向捲動快速讀完 inline 資料後選擇「進入 form 編輯」。",
             "steps": ["進入「產品管理 → 產品頁面」；清單可快速以品牌、機種、型號 查詢。",
                       "點「建立」建立新模板：填入「品牌 / 機種 / 型號 / OEM 編號」等主檔。",
                       "「上傳主圖」：點推薦圖上傳車型主圖；「車色」頁簽可透過「維護車色」彈窗設定品牌可供選車色與代碼。",
                       "「產品項」頁簽以「以車色為基準重建」按鈕，可依品牌車色表一鍵产生 SKU。"],
             "shots": []},
            {"h": "3.2 車型價格表（價目版本）",
             "brief": "以「價目版本」管理各車型的牌價、現金價與各期分期價；版本含生效日與狀態（草稿 / 生效 / 封存）。項目順序以 Phase 10 調整為：「機種」欄 sticky 凍結、標題列凍結；表格改為不換行、以橫向捲動避免揚壓；標題字體 12px；顧客贈品 / 附加費用說明 / 車色以「pre-line」完整換行顯示，不被裁切。\n\n每一列最左側提供「開啟詳細資料」外部連結按鈕，點選可跳到單筆產品的 form 畫面進行詳細編輯；inline 編輯仍不受影響。品牌 m2o 下拉不再被相鄰 sticky 欄遮擋。",
             "steps": ["進入「產品管理 → 價格表」檢視所有 SKU；可使用上方搜尋列以品牌 / 機種 / 型號 篩選。",
                       "inline 修改資料：點任一資料欄輸入內容後，点上方「儲存」按鈕即可。",
                       "詳細修改：點資料列最左侧的外部連結 icon 進入 form 畫面，適合大量欄位同時調整。",
                       "新增產品：請到「產品管理 → 產品頁面」建立模板後，於模板下頁簽「產品項」以「從此複製」或「新增」手動加入。"],
             "shots": [("21_pricetable_list", "車型價格表（sticky 機種欄，右上「開啟詳細資料」按鈕）"), ("21_pricetable_form", "價格表詳細 form 畫面")]},
            {"h": "3.3 價目版本與複製",
             "brief": "「價目版本」為價格表的現行生效集；可「由舊版複製」同步全部產品項后独立調整，未生效前主檔不受影響。",
             "steps": ["進入「產品管理 → 價目版本」點「建立」；選「源頭版本」可複製之前版本。",
                       "「生效」後原版本自動為「封存」，價格表與車銷單據都將讀取新生效版本。"],
             "shots": []},
            {"h": "3.4 分期規則模板 / 費用類型 / 規則掛接",
             "brief": "「分期規則模板」設定期數、首期月份、利率、最低順期。「費用類型」為車銷單可選的附加費用項目（如委託費、保険、職業謝金）主檔。「規則掛接」負責將費用類型與分期規則模板串接到某個價目版本，並可設定「只適用某些品牌 / 車型」。",
             "steps": ["「產品管理 → 分期規則模板」建立模板。",
                       "「產品管理 → 費用類型」建立代碼與預設金額。",
                       "「產品管理 → 規則掛接」選產品項 + 費用類型 + 分期規則模板複選掛接。"],
             "shots": [("26_installment_rule", "分期規則模板清單"), ("27_fee_type", "費用類型清單"), ("28_rule_binding", "規則掛接清單")]},
        ],
    },
    {
        "title": "4. 車銷管理（DMS Sale）",
        "intro": "車銷管理涵蓋車輛銷售單據、Excel 匯入與同步紀錄，是業績與佣金的計算來源。價格表與價目版本請見「3. 產品管理」。",
        "items": [
            {"h": "4.1 車輛銷售（Sale Order）",
             "brief": "合作車行成交時建立的車輛銷售單，是傭金與報表的核心紀錄。",
             "steps": ["進入「車銷管理 → 車輛銷售」。",
                       "點「建立」並選車行、車型、車架/引擎號、客戶資訊、售價。",
                       "售價可參考車型價格表手動填入；系統不自動帶入補助（補助由實際資料另行記錄）。確認後送出狀態流轉。"],
             "shots": [("22_sale_list", "車輛銷售清單"), ("22_sale_form", "車輛銷售表單")]},
            {"h": "4.2 同步紀錄（OrderProcessor Sync Log）",
             "brief": "記錄 OrderProcessor 將訂單資料匯入 DMIS 的成功 / 失敗 / 略過 / 已標記忽略狀態。\n\n讀取機制：優先剖析資料夾下的 result.json（同時支援舊格式 text_map / front / back 與新格式「以檔名為 key + docx text_content + 身分證 擷取欄位」）；若客戶姓名仍為空則 fallback 讀取同資料夾含「原始資料」工作表的 *.xlsx 補足。\n\n車行名稱比對在原字串找不到時，會 fallback 去除全的上下空白後重比（例「昌 億」→「昌億」）。\n\n以 folder_name 去重，log 有紀錄則不重複匯入。請在 form 頂部使用「重新同步」按鈕手動重試，或在 list view 勾選多筆後從「動作」選單選取「重新同步（重新解析備份資料夾）」以批次重試。",
             "steps": ["進入「車銷管理 → 同步紀錄」檢視最新同步結果。",
                       "點開單筆查看說明、folder_name、關聯訂單；需重試者點 form 頂部「重新同步」（會先刪原訂單與 log，再重新剖析、跳轉顯示新 log）。",
                       "批次重試：於 list view 勾選多筆→點上方「動作」→選「重新同步（重新解析備份資料夾）」。",
                       "歷史資料一次性忽略：勾選後點「動作 → 標記現有資料夾為歷史（不匯入）」。"],
             "shots": [("24_sync_log", "同步紀錄清單（含狀態色標與動作選單）"), ("24_sync_log_form", "同步紀錄 form 畫面（含「重新同步」按鈕）")]},
            {"h": "4.3 Excel 銷貨匯入",
             "brief": "支援以 Excel 檔上傳大批銷售單據，系統會驗證並建檔。",
             "steps": ["進入「車銷管理 → Excel 銷貨匯入」精靈。",
                       "下載樣板填寫後上傳；系統預覽資料行並回報錯誤列。",
                       "確認無誤後執行匯入，成功筆數會反映在車輛銷售清單。"],
             "shots": [("25_excel_import", "Excel 銷貨匯入精靈")]},
        ],
    },
    {
        "title": "5. 銷售分析（Metabase 儀表板）",
        "intro": "銷售分析統一以 Metabase 呈現，提供總車輛、油車 / 電車、車行、通路、客群等多維度儀表板，"
                 "使用者可在 http://localhost:3000 以 Metabase 帳號登入後查閱。"
                 "原 Odoo 內「銷售分析」選單下的原始數據查詢已隱藏（詳見 8.3 節）。",
        "items": [
            {"h": "5.1 進入 Metabase",
             "brief": "DMIS 的 BI 報表平台，與 Odoo 共用同一套 PostgreSQL，由部署腳本一鍵建立資料表與儀表板。",
             "steps": ["於瀏覽器開啟 http://localhost:3000 進入 Metabase 登入頁。",
                       "以 Metabase 管理員帳號登入後即可看到首頁，常用儀表板會置於 SUZUKI 銷售統計集合中。"],
             "shots": [("metabase_home", "Metabase 首頁（登入後）"),
                       ("metabase_dashboards", "儀表板清單（依集合分類）")]},
            {"h": "5.2 P1 總車輛銷售",
             "brief": "全車種（油車 + 電車）總體銷售儀表板：銷售來源 × 領牌年月、車種類型 × 領牌年月等。",
             "steps": ["於 Metabase 集合中點選「P1 總車輛銷售」。",
                       "可調整「領牌年月」與「銷售來源」篩選器後即時刷新圖表。"],
             "shots": [("metabase_p1_total_sales", "P1 總車輛銷售儀表板")]},
            {"h": "5.3 P3 電動車銷售統計",
             "brief": "專注電動車（含速克達 / 微型 / 重型 / 白牌）之銷售面向，與油車儀表板對照使用。",
             "steps": ["於集合中點選「P3 電動車銷售統計」。",
                       "可依時間、車型、通路切片，瞭解電車各車型佔比與走勢。"],
             "shots": [("metabase_p3_ev_sales", "P3 電動車銷售統計")]},
            {"h": "5.4 P9 油車銷售統計",
             "brief": "以油車為主的銷售儀表板，提供月別、車種、車型等銷售結構。",
             "steps": ["於集合中點選「P9 油車銷售統計」。"],
             "shots": [("metabase_p9_fuel_sales", "P9 油車銷售統計")]},
            {"h": "5.5 P11 油車-車行銷售統計",
             "brief": "以「合作車行」為主軸的油車銷售排行與貢獻度，便於評估車行表現。",
             "steps": ["於集合中點選「P11 油車-車行銷售統計」。",
                       "可使用搜尋列鎖定特定車行或時段檢視。"],
             "shots": [("metabase_p11_dealer_sales", "P11 油車-車行銷售統計")]},
            {"h": "5.6 其他內建儀表板",
             "brief": "系統另內建以下儀表板，可於 Metabase 集合中直接點開使用："
                      "P2 銷售機種統計、P4 基隆公益青年、P5 / P6 / P7 / P8 電動車系列、"
                      "P10 油車網路平台銷售、P12 油車佣金明細、P13 油車台數統計、"
                      "P14 ~ P19 客群 / 性別 / 顏色 / 區域 × 車型分析、P20 通路銷售、"
                      "P22 客群 × 車型分析等。",
             "steps": ["於 Metabase 集合中依命名前綴（P 編號）查找並點選即可。",
                       "如需新增報表，請聯絡系統管理員（資料模型由 scripts/metabase_*.py 部署）。"],
             "shots": []},
        ],
    },
    {
        "title": "6. 佣金管理（DMS Commission）",
        "intro": "佣金管理採「規則 → 結果 → 核銷」的三層模型，支援基礎佣金、車行/車種覆蓋、台數獎勵與多種激勵品。",
        "items": [
            {"h": "6.1 基礎傭金規則（依車型）",
             "brief": "依車型設定每台基礎傭金金額；為傭金計算的最底層規則。",
             "steps": ["進入「傭金管理 → 基礎傭金規則」。",
                       "建立規則並選擇車型、生效區間、傭金金額。",
                       "成交銷售單時系統會自動套用此規則。"],
             "shots": [("41_cm_product_rule", "基礎傭金規則清單"),
                       ("41_cm_product_rule_form", "基礎傭金規則表單")]},
            {"h": "6.2 車行覆蓋規則",
             "brief": "為特定合作車行客製傭金（疊加或取代基礎金額）。",
             "steps": ["進入「傭金管理 → 車行覆蓋規則」並指定車行 + 加碼/取代金額。"],
             "shots": [("42_cm_dealer_rule", "車行覆蓋規則清單")]},
            {"h": "6.3 車種覆蓋規則",
             "brief": "針對特定車型再加碼，例如新車促銷期間的單車型加碼。",
             "steps": ["進入「傭金管理 → 車種覆蓋規則」設定車型與加碼金額。"],
             "shots": [("43_cm_vehicle_rule", "車種覆蓋規則清單")]},
            {"h": "6.4 台數現金獎勵規則（Volume Rule）",
             "brief": "依當月達成台數階梯給予額外現金獎勵（如 ≥10 台加 1,000/台）。",
             "steps": ["進入「傭金管理 → 台數現金獎勵規則」。",
                       "建立階梯：起點台數、結束台數、單台/總額獎勵。"],
             "shots": [("44_cm_volume_rule", "台數現金獎勵規則清單")]},
            {"h": "6.5 台數實物獎勵規則",
             "brief": "達到指定台數時送出實物（如禮品、零件、提貨券）。",
             "steps": ["進入「傭金管理 → 台數實物獎勵規則」並指定品項與門檻。"],
             "shots": [("45_cm_volume_gift", "台數實物獎勵規則清單")]},
            {"h": "6.6 激勵品項定義",
             "brief": "建立可發放的激勵品（提貨券、零件、配件）主檔。",
             "steps": ["進入「傭金管理 → 激勵品項定義」並建立品項與單位。"],
             "shots": [("46_incentive_type", "激勵品項定義清單")]},
            {"h": "6.7 激勵觸發規則",
             "brief": "定義何種銷售條件觸發何種激勵品（連結銷售單與品項）。",
             "steps": ["進入「傭金管理 → 激勵觸發規則」設定觸發條件、品項、數量。"],
             "shots": [("47_incentive_rule", "激勵觸發規則清單")]},
            {"h": "6.8 傭金記錄（Commission Record）",
             "brief": "系統依規則自動產生的每筆傭金結果，可檢視/結案。",
             "steps": ["進入「傭金管理 → 傭金記錄」。",
                       "篩選車行/月份檢視結果；確認金額無誤後改為「結案」。"],
             "shots": [("48_cm_record", "傭金記錄清單"), ("48_cm_record_form", "傭金記錄表單")]},
            {"h": "6.9 激勵核銷（Incentive Delivery）",
             "brief": "管理實物激勵的給付狀態（待出貨/已出貨/作廢）。",
             "steps": ["進入「傭金管理 → 激勵核銷」。",
                       "選擇待出貨記錄，填入出貨日期、人員、追蹤號後狀態變更。"],
             "shots": [("49_incentive_delivery", "激勵核銷清單"),
                       ("49_incentive_delivery_form", "激勵核銷表單")]},
            {"h": "6.10 月結傭金報表",
             "brief": "依結案月份彙總所有傭金結果，作為對帳/出帳依據。",
             "steps": ["進入「傭金管理 → 月結傭金報表」即可瀏覽。"],
             "shots": [("4a_cm_monthly", "月結傭金報表")]},
            {"h": "6.11 傭金總報表",
             "brief": "跨期間檢視傭金總額並可依合作車行 / 業務切片。",
             "steps": ["進入「傭金管理 → 傭金總報表」。"],
             "shots": [("4b_cm_summary", "傭金總報表")]},
        ],
    },
    {
        "title": "7. 零件管理（DMS Parts）",
        "intro": "零件管理涵蓋零件主檔、分類、目錄與批次匯入工具，可同時支撐銷售贈品與激勵品的領用流程。",
        "items": [
            {"h": "7.1 零件清單",
             "brief": "管理所有零件之主檔（品號、名稱、適用車型、單價、庫存）。",
             "steps": ["進入「零件管理 → 零件清單」。",
                       "點「建立」輸入品號、名稱、適用車型、單位、單價後儲存。",
                       "可在搜尋列以品號或名稱快速定位。"],
             "shots": [("51_part_list", "零件清單"), ("51_part_form", "零件表單")]},
            {"h": "7.2 零件分類",
             "brief": "分類維護（傳動、剎車、燈具、保養品 …）便於統計與目錄整理。",
             "steps": ["進入「零件管理 → 零件分類」並建立樹狀分類。"],
             "shots": [("52_part_category", "零件分類")]},
            {"h": "7.3 零件目錄（Catalog）",
             "brief": "依品牌/車型編製零件目錄，支援版本管理與發佈。",
             "steps": ["進入「零件管理 → 零件目錄」。",
                       "建立或開啟目錄；新增明細並關聯零件主檔。"],
             "shots": [("53_catalog_list", "零件目錄清單"), ("53_catalog_form", "零件目錄表單")]},
            {"h": "7.4 零件查詢工具",
             "brief": "供前線人員快速查詢適用零件之精靈。",
             "steps": ["進入「零件管理 → 零件查詢」精靈，輸入車型/年份即顯示適用品。"],
             "shots": [("54_catalog_search_wizard", "零件查詢工具")]},
            {"h": "7.5 CSV 批次匯入",
             "brief": "提供中文表頭、含 UTF-8 BOM 的標準 CSV 樣板（Excel 雙擊不亂碼），一次大量新增/更新零件與目錄明細。",
             "steps": ["進入「零件管理 → CSV 批次匯入」精靈。",
                       "點「下載樣板」按鈕取得『零件目錄匯入樣板.csv』，表頭為：目錄名稱 / 分區代碼 / 分區名稱 / 分區類別 / 序號 / 零件編號 / 零件名稱 / 單位 / 數量 / 建議售價。",
                       "填寫完成後上傳檔案；系統會驗證每筆資料格式並回報。表頭可採中文或英文（catalog_name, section_code…）。",
                       "確認無誤後執行匯入，新增/更新筆數於完成頁顯示。"],
             "shots": [("55_catalog_import_wizard", "CSV 批次匯入精靈（內含「下載樣板」按鈕）")]},
        ],
    },
    {
        "title": "8. 使用者管理系統（User Management）",
        "intro": "使用者管理系統負責功能權限分組以及全系統的操作稽核，是合規與資安的關鍵環節。",
        "items": [
            {"h": "8.1 存取群組管理",
             "brief": "依角色（總部、區主管、業務、合作車行）建立功能/資料權限群組。",
             "steps": ["進入「使用者管理系統 → 存取群組管理」。",
                       "建立群組並勾選可用模組與資料範圍。",
                       "在 Odoo 使用者主檔將其加入相應群組生效。"],
             "shots": [("61_access_group_list", "存取群組清單"),
                       ("61_access_group_form", "存取群組表單")]},
            {"h": "8.2 操作歷程（Audit Log）",
             "brief": "完整記錄系統內 CRUD 動作，包含操作人、時間、模型、舊/新值。",
             "steps": ["進入「使用者管理系統 → 操作歷程」。",
                       "可使用搜尋列鎖定使用者、模型或時段；點開單筆檢視差異內容。"],
             "shots": [("62_audit_log_list", "操作歷程清單"),
                       ("62_audit_log_form", "操作歷程詳細")]},
            {"h": "8.3 介面精簡（隱藏未使用之根選單與仅管理員可見項目）",
             "brief": "為簡化日常操作，已透過 user_management 模組將以下選單隱藏："
                      "「討論」「財務結算」「庫存」三個根選單，以及「銷售分析 → 原始數據查詢」子選單。另外「車行管理 → 品牌」選單仅系統管理員（base.group_system）可見，避免一般使用者誤改品牌主檔。",
             "steps": ["隅藏設定統一寫於 addons/user_management/data/hide_menus.xml；",
                       "品牌選單的可見性寫於 addons/dms_core/views/brand_views.xml 之 menu_brands.groups；",
                       "若日後需重新顯示，將對應 record 的 active 改為 True 並升級模組即可；",
                       "此調整僅影響選單可見性，不刪除任何資料或模組功能。"],
             "shots": []},
        ],
    },
]


def section_module(doc, sec):
    add_heading(doc, sec["title"], level=1)
    add_para(doc, sec["intro"], size=11)
    for item in sec["items"]:
        add_heading(doc, item["h"], level=2)
        add_para(doc, item["brief"], size=11, color=GRAY)
        add_heading(doc, "操作步驟", level=3)
        add_steps(doc, item["steps"])
        add_heading(doc, "畫面截圖", level=3)
        for n, cap in item["shots"]:
            add_image(doc, n, cap)
    doc.add_page_break()


def section_validation(doc):
    add_heading(doc, "9. 驗證與重現步驟", level=1)
    add_para(doc, "使用者可依下列指令於本機重現此報告所有畫面與資料來源：")
    add_kv_table(doc, [
        ("啟動服務", "make up"),
        ("健康檢查", "make smoke 或 bash scripts/smoke_odoo.sh"),
        ("檢視容器", "docker compose ps"),
        ("登入網址", "http://localhost:8069  (帳號 admin / 密碼 admin)"),
        ("重新產報告", "python3 scripts/seed_temp_for_report.py seed → "
                       "python3 scripts/capture_ui_full.py → "
                       "python3 scripts/build_progress_report.py → "
                       "python3 scripts/seed_temp_for_report.py cleanup"),
    ])

    add_heading(doc, "9. 結論", level=1)
    add_para(doc,
             "DMIS 已完成「合作車行 → 銷售 → 傭金 → 零件 → 報表 → 使用者」六大主軸的端到端能力，"
             "並完整支援 Excel/CSV 大量匯入、多層傭金規則、操作稽核，以及 Metabase BI 整合。"
             "後續將持續優化效能、報表自動化與與外部 OrderProcessor 的雙向同步。")


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    cover(doc)
    section_overview(doc)
    for s in SECTIONS:
        section_module(doc, s)
    section_validation(doc)

    doc.save(DOCX)
    print(f"DONE -> {DOCX}")


if __name__ == "__main__":
    main()
