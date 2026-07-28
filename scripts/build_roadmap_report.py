#!/usr/bin/env python3
"""產生 DMIS《後續優化與功能擴充規劃》Word 文件。

風格與 build_progress_report.py 一致（Microsoft JhengHei + 海軍藍標題）。
不依賴截圖，純文字 + 表格，便於與進度報告並列交付。
"""
import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT_DIR = Path("/home/audi/project/DMIS/output_report")
DOCX = OUT_DIR / "DMIS_後續優化與擴充規劃_v1.docx"

ZH = "Microsoft JhengHei"
NAVY = RGBColor(0x1F, 0x3A, 0x68)
GRAY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)


def set_zh(run, size=11, bold=False, color=None):
    run.font.name = ZH
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), ZH)
    rFonts.set(qn("w:hAnsi"), ZH)
    rFonts.set(qn("w:ascii"), ZH)
    run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, size=11, bold=False, color=None, align=None, indent=None,
             space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_zh(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 26, 1: 20, 2: 15, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_zh(run, size=sizes.get(level, 12), bold=True, color=NAVY)


def add_bullets(doc, items, indent=0.2):
    for s in items:
        add_para(doc, f"• {s}", size=11, indent=indent, space_after=2)


def add_priority_table(doc, rows):
    """rows = [(項目, 內容, 預期效益)]"""
    t = doc.add_table(rows=len(rows) + 1, cols=3)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["項目", "內容說明", "預期效益"]
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_zh(cell.paragraphs[0].add_run(h), size=10, bold=True, color=NAVY)
    widths = [Inches(1.4), Inches(3.6), Inches(1.6)]
    for r, row in enumerate(rows, 1):
        for i, val in enumerate(row):
            cell = t.rows[r].cells[i]
            cell.width = widths[i]
            set_zh(cell.paragraphs[0].add_run(val), size=10)


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "DMIS 合作車行管理系統", size=30, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "後續優化與功能擴充規劃（v1）", size=18, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "技術優化｜功能加強｜標桿系統參考｜BI 與 AI 強化", size=13, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(5):
        doc.add_paragraph()
    add_para(doc, f"撰寫日期：{datetime.date.today().isoformat()}", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
    add_para(doc, "技術基礎：Odoo 16 Community｜PostgreSQL 15｜Docker Compose｜Metabase",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
    add_para(doc, "適用對象：專案管理者、IT 主管、營運／業務主管、外部審查",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
    doc.add_page_break()


def section_intro(doc):
    add_heading(doc, "1. 文件目的", level=1)
    add_para(doc,
             "本文件接續《DMIS 操作手冊與進度報告（v5）》，"
             "彙整現階段已交付的六大模組（車行、車銷、銷售分析、傭金、零件、使用者）"
             "之後續可優化方向與待擴充功能，並參考國內外類似 DMS／經銷商管理系統的常見能力，"
             "提供決策層作為下一階段規劃的依據。本文件不替代既有 specs/ 規格，"
             "凡列入排程之項目，仍須先於 specs/ 下建立或更新對應規格再進行開發。", size=11)

    add_heading(doc, "2. 規劃原則", level=1)
    add_bullets(doc, [
        "以「最小可用 → 數據驅動 → 智能輔助」三階段心智模型推進，先補齊基礎，再導入分析與 AI。",
        "凍結模組（dms_core）採非侵入式擴充：以 _inherit 或新建獨立模組實作，避免影響既有資料。",
        "任何新增欄位／資料結構變更，皆以 spec-first（specs/）→ migration → 自動化測試的方式落地。",
        "外部整合（OEM、會計、金流）以 API 中介層解耦，避免 Odoo 直接耦合外部系統。",
        "所有後續報表與儀表板優先以 Metabase 視覺化呈現，Odoo 端僅保留資料維護與快速查詢。",
    ])


def section_tech(doc):
    add_heading(doc, "3. 技術面優化", level=1)

    add_heading(doc, "3.1 效能與資料庫", level=2)
    add_priority_table(doc, [
        ("PostgreSQL 索引盤點",
         "針對 dms.sale、dms.commission.record、dms.visit 等高頻查詢資料表建立複合索引，"
         "並導入 pg_stat_statements 監控慢查詢。",
         "查詢延遲下降 50% 以上"),
        ("資料分區（Partitioning）",
         "對 dms.sale、dms.audit.log、dms.report.* 依年月分區，"
         "歷史資料可獨立歸檔避免拖累線上查詢。",
         "降低單表大小、提升維運彈性"),
        ("Materialized View",
         "銷售分析、傭金月報以 PostgreSQL Materialized View 預計算，"
         "搭配 cron 定時 refresh，避免 pivot 即時計算負載。",
         "報表載入由分鐘級降至秒級"),
        ("Redis 快取層",
         "Odoo session、ir.attachment 縮圖、Metabase 查詢結果接入 Redis，"
         "降低資料庫與檔案系統壓力。",
         "Web 互動更順暢"),
        ("Excel 匯入優化",
         "現行匯入採逐筆 create，改為批次 insert + 暫存表 + 校驗報告，"
         "支援大量資料（>10 萬筆）匯入。",
         "匯入時間下降 70%＋"),
    ])

    add_heading(doc, "3.2 架構與模組化", level=2)
    add_priority_table(doc, [
        ("API 中介層",
         "於 Odoo 之上建立獨立 API Gateway（FastAPI / Nginx + JWT），"
         "對外（OEM、行動 App、第三方）統一介接點，避免直接暴露 Odoo RPC。",
         "資安提升、整合彈性"),
        ("事件驅動（Outbox Pattern）",
         "重要業務事件（成交、傭金結算、訪店完成）寫入 outbox table，"
         "再由 worker 投遞至訊息佇列（Redis Streams／RabbitMQ）。",
         "為跨系統整合與審計鋪路"),
        ("計算邏輯抽離為服務",
         "傭金、報表規則、虛擬欄位等計算邏輯可逐步抽離為獨立 Python 服務，"
         "Odoo 僅負責資料維護與顯示。",
         "降低 Odoo 升級耦合"),
        ("多租戶／多公司強化",
         "若未來需服務多家經銷集團，啟用 Odoo multi-company 並補強 record rule 與資料隔離測試。",
         "支援 SaaS 化營運"),
    ])

    add_heading(doc, "3.3 部署、CI/CD 與監控", level=2)
    add_priority_table(doc, [
        ("環境分層",
         "建立 dev／staging／prod 三層 docker compose，prod 啟用 read-replica DB 與每日邏輯備份。",
         "降低正式環境風險"),
        ("CI/CD 自動化",
         "GitHub Actions：lint（pylint-odoo）→ unit test（make smoke）→ "
         "build image → deploy to staging → 手動 approve → prod。",
         "縮短交付週期、降低人為錯誤"),
        ("觀測性（Observability）",
         "整合 Prometheus + Grafana 監控 Odoo workers、PostgreSQL、Metabase；"
         "Loki 收集應用 log；Sentry 捕捉錯誤。",
         "故障 5 分鐘內可定位"),
        ("災難復原（DR）",
         "PostgreSQL WAL-G 異地備份、Metabase 設定備份，"
         "建立 RPO ≤ 15 分鐘、RTO ≤ 1 小時的還原演練手冊。",
         "符合企業級營運要求"),
        ("Cloudflare Tunnel 強化",
         "現行 cloudflared 可加上 Access Policy（OTP／SSO），"
         "限制 Odoo / Metabase 僅授權 IP 或員工 SSO 登入。",
         "對外存取更安全"),
    ])

    add_heading(doc, "3.4 資安與合規", level=2)
    add_priority_table(doc, [
        ("OWASP Top 10 自動掃描",
         "每次 PR 觸發 SAST（bandit）與相依套件掃描（pip-audit、Trivy on docker image）。",
         "預防已知弱點"),
        ("敏感資料遮罩",
         "客戶身分證、聯絡電話於非業務必要欄位顯示時自動遮罩，匯出 Excel 也須套用。",
         "符合個資保護要求"),
        ("操作稽核強化",
         "user_management.audit.log 補上『資料變更前後』欄位，"
         "敏感欄位（價格、傭金、客戶）強制留紀錄。",
         "可追溯性提升"),
        ("MFA／SSO",
         "Odoo 接 auth_oauth + Google / Azure AD，並對管理者群組強制啟用 MFA。",
         "帳號被盜風險降低"),
        ("最小權限稽核",
         "每季自動產出 ACL／record rule 差異報告，避免權限蔓延。",
         "符合稽核要求"),
    ])


def section_features(doc):
    add_heading(doc, "4. 功能面加強（依模組）", level=1)

    add_heading(doc, "4.1 車行管理（dms_core / dms_visit）", level=2)
    add_bullets(doc, [
        "車行健檢儀表板：合作門市的成交、傭金、回款、訪店頻率彙整為單一頁面，異常自動標紅。",
        "訪店行程最佳化：依地理位置 + 上次訪店日 + 業務日曆，自動排出最佳路線（TSP 演算法）。",
        "拜訪相片／簽到 GPS：行動端拍照即時上傳，含經緯度水印，避免造假。",
        "車行合約／附件管理：合約到期提醒、文件版本歷程、線上電子簽章（整合 DocuSign / 國發會 TWID）。",
        "車行分級制度：A/B/C 級依績效自動評等，連動傭金與供貨優先順序。",
    ])

    add_heading(doc, "4.2 車銷管理（dms_sale / dms_product）", level=2)
    add_bullets(doc, [
        "報價／訂單／交車三段式流程：可由報價單一鍵轉訂單、訂單一鍵轉交車單，狀態同步傭金。",
        "車型組態器（Configurator）：選色、選配、貸款方案於同一畫面試算，輸出 PDF 報價。",
        "庫存／配車模組：與 OEM 出貨資料介接，可看 VIN 階段、預計到車日、配車排程。",
        "二手車收購／媒合：估價工具、收購車況表、與新車交易連動以舊換新。",
        "信用查詢／徵信整合：與聯徵中心或第三方信用 API 介接，提早識別風險客戶。",
        "客製化報價樣板：依品牌、車系、活動切換報價單版面與條款。",
    ])

    add_heading(doc, "4.3 銷售分析（dms_report / dms_report_rule / dms_report_virtual）",
                level=2)
    add_bullets(doc, [
        "銷售漏斗分析：來客 → 試車 → 報價 → 成交 各階段轉換率與失單原因。",
        "車型／顏色／配備熱度：以圖表呈現各區域、各門市熱銷組合，回饋給 OEM 規劃。",
        "預測報表：以時間序列模型（Prophet / ARIMA）對下季銷售給出區間預測。",
        "規則引擎可視化：dms_report_rule 增加拖拉式條件編輯器，業務無須寫公式即可建規則。",
        "虛擬欄位 SQL 預覽：dms_report_virtual 支援即時試算與錯誤提示。",
    ])

    add_heading(doc, "4.4 傭金管理（dms_commission）", level=2)
    add_bullets(doc, [
        "傭金試算器：業務／主管於成交前可即時試算個人 + 團隊獎金。",
        "多階獎金結構：支援個人、組長、店長、區經理四階分潤與級距獎金（target bonus）。",
        "傭金結算工作流：月結 → 主管覆核 → 財務確認 → 匯款檔，整段流程留軌跡。",
        "獎金說明書 PDF：每期自動產出個人獎金明細書，可直接 email 發送。",
        "傭金爭議申訴單：業務可線上申訴，主管端有處理時效 SLA。",
    ])

    add_heading(doc, "4.5 零件管理（dms_parts）", level=2)
    add_bullets(doc, [
        "零件庫存即時化：與門市 POS／倉儲整合，安全庫存自動補貨建議。",
        "圖型零件查詢（EPC）：保留現有 catalog，加上點選爆炸圖直接帶料號的互動體驗。",
        "通用件交叉表：跨車型／品牌的相容料號自動建議。",
        "工時／工單系統雛形：保養工單對應零件耗用，連動到客戶車輛保養履歷。",
        "供應商比價／採購單：多供應商報價比較、自動產生採購單。",
    ])

    add_heading(doc, "4.6 使用者管理（user_management）", level=2)
    add_bullets(doc, [
        "群組權限模板：常見職務（業務、店長、財務、總部 BI）一鍵套用權限。",
        "離職／調動工作流：HR 通知 → 自動停權 → 主管確認 → 留檔。",
        "App／Web 雙裝置 session 管理：可遠端登出、查詢登入歷史。",
        "操作行為熱力圖：分析常用功能與冷門功能，回饋 UI 改版優先序。",
    ])


def section_ux(doc):
    add_heading(doc, "5. 使用者體驗、行動裝置與離線", level=1)
    add_bullets(doc, [
        "行動端業務 App：以 Odoo Mobile / Flutter 自開，覆蓋訪店、報價、簽到、相片上傳、傭金查詢。",
        "離線優先（Offline-first）：訪店與報價支援離線暫存、上線自動同步，山區與地下室可用。",
        "語音／OCR 輸入：拍身分證、行照即可帶入客戶與車輛資料；語音記錄訪店摘要。",
        "通知中心：交車、合約到期、傭金入帳等以 Web Push / LINE Notify / Email 多通道。",
        "深色模式與大字級：適合長時間使用與年長業務。",
        "i18n 強化：完整繁中／英文／東南亞語系（如越南、印尼），便於跨國經銷集團使用。",
    ])


def section_integration(doc):
    add_heading(doc, "6. 與現有 ERP／會計系統整合", level=1)
    add_bullets(doc, [
        "標準化交換格式：以 JSON Schema / OpenAPI 定義雙向介接介面，文件版控於 specs/。",
        "ERP 對接（鼎新、SAP、Oracle）：客戶、訂單、發票、應收/應付，採當日批次 + 即時補單。",
        "電子發票：串接財政部 B2C / B2B 電子發票 API，發票自動開立、折讓、作廢。",
        "會計傳票自動化：成交、退車、傭金核發自動產生分錄，會計僅做覆核。",
        "金流／第三方支付：信用卡、ATM、行動支付（街口、Line Pay）回饋至訂單付款狀態。",
        "OEM／原廠資料：與 Yamaha／其他原廠 EDI／API 對接配車、保固、召回、零件目錄更新。",
        "CRM / 客服：與 LINE OA、客服中心整合，客戶線上預約進廠、線上估價即進入車銷漏斗。",
        "HR / 出勤：業務出勤、訪店打卡資料雙向同步，連動傭金結算之請假計算。",
    ])


def section_bi(doc):
    add_heading(doc, "7. 資料分析、BI 與 AI 強化", level=1)

    add_heading(doc, "7.1 Metabase 儀表板擴充", level=2)
    add_bullets(doc, [
        "高層儀表板：總部一頁式儀表板（KPI、地圖、Top10 車行、月趨勢），支援大螢幕牆。",
        "區業／店長儀表板：依登入者身份自動過濾，無需切換角色。",
        "OEM 視角儀表板：依品牌／車型整合銷售、保固、零件耗用，作為與原廠對話依據。",
    ])

    add_heading(doc, "7.2 資料倉儲（DW）與湖倉", level=2)
    add_bullets(doc, [
        "建立 dms_dw：以 dbt / Airbyte 將 Odoo OLTP 資料定時匯入 PostgreSQL DW（或 ClickHouse）。",
        "Star Schema 模型：fact_sale、fact_commission、fact_visit + dim_dealer/dim_product/dim_time。",
        "Metabase 與 BI 工具直接查 DW，避免影響線上服務。",
    ])

    add_heading(doc, "7.3 AI／機器學習應用", level=2)
    add_priority_table(doc, [
        ("銷售預測",
         "依車型 × 區域 × 月份建立時間序列模型，輸出未來 3 個月預測區間。",
         "備車與促銷更精準"),
        ("流失客戶預測",
         "以車輛保養／保險／聯絡頻率特徵建立分類模型，識別高流失風險客戶。",
         "保留率提升"),
        ("最佳促銷組合",
         "以歷史成交資料做 Uplift Modeling，找出對成交率影響最大的折讓／贈品組合。",
         "活動投報率提升"),
        ("文件智能化（LLM）",
         "合約／報價條款比對、客服 FAQ、業務話術建議由內部 LLM（self-hosted）提供。",
         "降低重複問答工時"),
        ("影像辨識",
         "車牌、VIN、行照、駕照辨識自動帶入欄位，並偵測訪店相片造假。",
         "資料品質提升"),
    ])


def section_benchmark(doc):
    add_heading(doc, "8. 標桿與類似系統參考", level=1)
    add_para(doc,
             "下列系統為國內外常見的 DMS／經銷商管理 SaaS，列出可借鏡的功能設計。"
             "本表僅作為功能對照參考，不代表計畫直接整合或採購。", size=11, color=GRAY)

    add_heading(doc, "8.1 國際 DMS／經銷商平台", level=2)
    add_priority_table(doc, [
        ("CDK Global",
         "美國最大 DMS 之一，涵蓋新車、二手車、零件、保養、會計、CRM 全模組。",
         "整體模組分工可參考"),
        ("Reynolds & Reynolds",
         "強項在會計傳票、合約電子化與 OEM 介接，提供 ERA-IGNITE 平台。",
         "會計整合與 OEM 介接"),
        ("Dealertrack DMS（Cox Automotive）",
         "雲端架構、開放 API、報表彈性高，與第三方 CRM／金流整合度佳。",
         "雲端與開放 API 思維"),
        ("Tekion ARC",
         "新一代 cloud-native DMS，主打單一資料平台 + AI 推薦。",
         "AI 與單一資料平台"),
        ("Autoline (Keyloop)",
         "歐洲廣泛採用，多語系、多公司、多幣別支援完整。",
         "多公司與多語系設計"),
    ])

    add_heading(doc, "8.2 中文／亞洲區常見系統", level=2)
    add_priority_table(doc, [
        ("中華汽車 e-DMS",
         "以原廠端為核心，強化保固、召回、零件分發與經銷商評鑑。",
         "OEM 與經銷商雙向流程"),
        ("Yamaha／三陽 經銷管理系統",
         "強調機車產業 SKU、零件目錄、機車保養工時。",
         "機車零件與保養工時"),
        ("和泰 TOYOTA T-Connect",
         "車主 App、保養預約、車況遠端監控與 DMS 串聯。",
         "車主端 App 與預約"),
        ("中國 4S 店常見系統（ePlus、車智匯）",
         "強化線上引流、線索分配、試駕到店追蹤、二手車置換。",
         "線索分配與漏斗"),
    ])

    add_heading(doc, "8.3 開源／可自託管選項", level=2)
    add_bullets(doc, [
        "Odoo Enterprise 模組：CRM、Sign、Studio、Documents 可考慮選購補強。",
        "ERPNext：模組完整、開源、有汽車經銷示範資料，可作為對照。",
        "Frappe Insights / Apache Superset：除 Metabase 外的 BI 替代方案。",
        "Apache Airflow / Dagster：用於資料倉儲與報表的 ETL 排程。",
        "MinIO / S3：作為合約、相片、報表 PDF 的物件儲存後端。",
    ])


def section_governance(doc):
    add_heading(doc, "9. 治理與後續行動建議", level=1)
    add_bullets(doc, [
        "成立月度 Roadmap 評審：由 PM、IT、業務、財務代表共同決定優先序，產出更新版本文件。",
        "每項規劃落地前先建立 specs/NNN-<topic>/ 五份規格（charter、scope、design、test、release）。",
        "對外整合一律先簽 NDA／API 規格書，避免實作後再回頭談協定。",
        "AI／資料分析項目應先設定明確的成功指標（如轉換率提升 X%、客單價提升 Y 元）。",
        "持續維護 docs/CHANGELOG.md 與 docs/USER_MANUAL.md，作為對使用者的單一真相來源。",
    ])

    add_heading(doc, "10. 結語", level=1)
    add_para(doc,
             "DMIS 在現階段已完成六大模組的基礎建置，具備可用於日常營運的完整能力。"
             "後續若能依本文件之優先序，逐步補上效能、整合、BI 與 AI 能力，"
             "將可讓 DMIS 從『內部營運工具』進化為『經銷體系數位中樞』，"
             "支撐總部與合作車行在快速變化的汽機車市場中持續競爭。",
             size=11)
    add_para(doc, "—— 本文件僅供規劃與內部審查使用，實際開發以 specs/ 規格書為準 ——",
             size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    # 預設樣式統一中文字型
    style = doc.styles["Normal"]
    style.font.name = ZH
    style.font.size = Pt(11)

    cover(doc)
    section_intro(doc)
    section_tech(doc)
    section_features(doc)
    section_ux(doc)
    section_integration(doc)
    section_bi(doc)
    section_benchmark(doc)
    section_governance(doc)

    doc.save(DOCX)
    print(f"DONE -> {DOCX}")


if __name__ == "__main__":
    main()
